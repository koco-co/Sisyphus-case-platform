import io

from docx import Document


def parse_docx(raw_bytes: bytes) -> tuple[str, dict]:
    doc = Document(io.BytesIO(raw_bytes))
    sections: list[dict[str, str]] = []
    current_heading = ""
    current_body_lines: list[str] = []
    full_text_lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        full_text_lines.append(text)

        if para.style and para.style.name and para.style.name.startswith("Heading"):
            if current_heading or current_body_lines:
                sections.append(
                    {
                        "heading": current_heading,
                        "body": "\n".join(current_body_lines).strip(),
                    }
                )
            current_heading = text
            current_body_lines = []
        else:
            current_body_lines.append(text)

    if current_heading or current_body_lines:
        sections.append(
            {
                "heading": current_heading,
                "body": "\n".join(current_body_lines).strip(),
            }
        )

    # Extract tables
    tables_data: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables_data.append(rows)

    full_text = "\n".join(full_text_lines)
    content_ast: dict = {
        "raw_text": full_text,
        "sections": sections,
        "tables": tables_data,
    }

    return full_text, content_ast
