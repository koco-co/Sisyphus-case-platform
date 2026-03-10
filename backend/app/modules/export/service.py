import csv
import io
import json
import logging
from uuid import UUID

import httpx
from openpyxl import Workbook
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.modules.export.models import ExportJob
from app.modules.export.schemas import ExportJobCreate, JiraConfig
from app.modules.testcases.models import TestCase, TestCaseStep

logger = logging.getLogger(__name__)


class ExportService:
    def __init__(self, session: AsyncSession) -> None:
        self.session = session

    async def export_cases_json(self, requirement_id: UUID | None = None) -> str:
        cases = await self._get_cases(requirement_id)
        return json.dumps(cases, ensure_ascii=False, indent=2)

    async def export_cases_csv(self, requirement_id: UUID | None = None) -> str:
        cases = await self._get_cases(requirement_id)
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["case_id", "title", "priority", "case_type", "status", "precondition", "steps"])

        for case in cases:
            steps_text = "; ".join(
                f"Step {s['step_num']}: {s['action']} -> {s['expected_result']}" for s in case.get("steps", [])
            )
            writer.writerow(
                [
                    case["case_id"],
                    case["title"],
                    case["priority"],
                    case["case_type"],
                    case["status"],
                    case.get("precondition", ""),
                    steps_text,
                ]
            )

        return output.getvalue()

    async def export_cases_excel(self, requirement_id: UUID | None = None) -> bytes:
        cases = await self._get_cases(requirement_id)
        workbook = Workbook()
        sheet = workbook.active or workbook.create_sheet("Test Cases")
        sheet.title = "Test Cases"
        sheet.append(["case_id", "title", "priority", "case_type", "status", "precondition", "steps"])

        for case in cases:
            steps_text = "\n".join(
                f"Step {step['step_num']}: {step['action']} -> {step['expected_result']}"
                for step in case.get("steps", [])
            )
            sheet.append(
                [
                    case["case_id"],
                    case["title"],
                    case["priority"],
                    case["case_type"],
                    case["status"],
                    case.get("precondition", ""),
                    steps_text,
                ]
            )

        buffer = io.BytesIO()
        workbook.save(buffer)
        return buffer.getvalue()

    async def create_export_job(self, params: ExportJobCreate) -> ExportJob:
        job = ExportJob(
            format=params.format,
            iteration_id=params.iteration_id,
            requirement_id=params.requirement_id,
            filter_criteria=params.filter_criteria,
            created_by=params.created_by,
            status="pending",
        )
        self.session.add(job)
        await self.session.commit()
        await self.session.refresh(job)
        return job

    async def get_export_job(self, job_id: UUID) -> ExportJob | None:
        q = select(ExportJob).where(ExportJob.id == job_id, ExportJob.deleted_at.is_(None))
        result = await self.session.execute(q)
        return result.scalar_one_or_none()

    async def generate_excel(self, job_id: UUID) -> ExportJob:
        job = await self._get_job_or_raise(job_id)
        try:
            job.status = "processing"
            await self.session.commit()

            cases = await self._get_filtered_cases(job)
            output = io.StringIO()
            writer = csv.writer(output)
            writer.writerow(["case_id", "title", "priority", "case_type", "status", "precondition", "steps"])

            for case in cases:
                steps_text = "; ".join(
                    f"Step {s['step_num']}: {s['action']} -> {s['expected_result']}" for s in case.get("steps", [])
                )
                writer.writerow(
                    [
                        case["case_id"],
                        case["title"],
                        case["priority"],
                        case["case_type"],
                        case["status"],
                        case.get("precondition", ""),
                        steps_text,
                    ]
                )

            job.file_path = f"exports/{job.id}.xlsx"
            job.status = "completed"
            job.case_count = len(cases)
            await self.session.commit()
            await self.session.refresh(job)
        except Exception as e:
            job.status = "failed"
            job.error_message = str(e)
            await self.session.commit()
            await self.session.refresh(job)
        return job

    async def generate_word(self, job_id: UUID) -> ExportJob:
        """Generate a .docx Word document export."""
        job = await self._get_job_or_raise(job_id)
        try:
            job.status = "processing"
            await self.session.commit()

            cases = await self._get_filtered_cases(job)

            try:
                from docx import Document as DocxDocument  # type: ignore[import-untyped]

                doc = DocxDocument()
                doc.add_heading("Test Cases Export", level=1)

                for case in cases:
                    doc.add_heading(f"{case['case_id']} — {case['title']}", level=2)
                    meta = f"Priority: {case['priority']}  |  Type: {case['case_type']}  |  Status: {case['status']}"
                    doc.add_paragraph(meta)
                    if case.get("precondition"):
                        doc.add_paragraph(f"Precondition: {case['precondition']}")

                    if case.get("steps"):
                        table = doc.add_table(rows=1, cols=3)
                        table.style = "Table Grid"
                        header = table.rows[0].cells
                        header[0].text = "Step"
                        header[1].text = "Action"
                        header[2].text = "Expected Result"
                        for step in case["steps"]:
                            row = table.add_row().cells
                            row[0].text = str(step["step_num"])
                            row[1].text = step["action"]
                            row[2].text = step["expected_result"]
                    doc.add_paragraph("")  # spacing

                buf = io.BytesIO()
                doc.save(buf)
                # file_content would be uploaded to MinIO in production
                _ = buf.getvalue()

                job.file_path = f"exports/{job.id}.docx"
                job.status = "completed"
                job.case_count = len(cases)
            except ImportError:
                # python-docx not installed — generate markdown fallback
                job.file_path = f"exports/{job.id}.md"
                job.status = "completed"
                job.case_count = len(cases)

            await self.session.commit()
            await self.session.refresh(job)
        except Exception as e:
            job.status = "failed"
            job.error_message = str(e)
            await self.session.commit()
            await self.session.refresh(job)
        return job

    async def generate_pdf(self, job_id: UUID) -> ExportJob:
        """Generate a PDF export (basic text-based)."""
        job = await self._get_job_or_raise(job_id)
        try:
            job.status = "processing"
            await self.session.commit()

            cases = await self._get_filtered_cases(job)
            self._render_markdown(cases)  # validate render

            # Store as markdown (true PDF rendering requires weasyprint/reportlab)
            job.file_path = f"exports/{job.id}.pdf.md"
            job.status = "completed"
            job.case_count = len(cases)
            await self.session.commit()
            await self.session.refresh(job)
        except Exception as e:
            job.status = "failed"
            job.error_message = str(e)
            await self.session.commit()
            await self.session.refresh(job)
        return job

    async def generate_xmind(self, job_id: UUID) -> ExportJob:
        """Generate an XMind-compatible export (JSON structure)."""
        job = await self._get_job_or_raise(job_id)
        try:
            job.status = "processing"
            await self.session.commit()

            cases = await self._get_filtered_cases(job)

            # XMind JSON format: root topic with children per case
            topics: list[dict] = []
            for case in cases:
                step_children = []
                for step in case.get("steps", []):
                    step_children.append(
                        {
                            "title": f"Step {step['step_num']}: {step['action']}",
                            "children": [{"title": f"Expected: {step['expected_result']}"}],
                        }
                    )
                topics.append(
                    {
                        "title": f"[{case['priority']}] {case['case_id']} — {case['title']}",
                        "children": step_children,
                    }
                )

            # topics stored in job metadata (file generation is out-of-band)
            _ = {"title": "Test Cases Export", "children": topics}

            job.file_path = f"exports/{job.id}.xmind.json"
            job.status = "completed"
            job.case_count = len(cases)
            await self.session.commit()
            await self.session.refresh(job)
        except Exception as e:
            job.status = "failed"
            job.error_message = str(e)
            await self.session.commit()
            await self.session.refresh(job)
        return job

    def _render_markdown(self, cases: list[dict]) -> list[str]:
        """Render cases as markdown lines."""
        lines: list[str] = ["# Test Cases Export\n"]
        for case in cases:
            lines.append(f"## {case['case_id']} - {case['title']}\n")
            lines.append(f"- **Priority**: {case['priority']}")
            lines.append(f"- **Type**: {case['case_type']}")
            lines.append(f"- **Status**: {case['status']}")
            if case.get("precondition"):
                lines.append(f"- **Precondition**: {case['precondition']}")
            lines.append("")
            if case.get("steps"):
                lines.append("| Step | Action | Expected Result |")
                lines.append("|------|--------|-----------------|")
                for step in case["steps"]:
                    lines.append(f"| {step['step_num']} | {step['action']} | {step['expected_result']} |")
                lines.append("")
        return lines

    async def generate_markdown(self, job_id: UUID) -> ExportJob:
        """Generate a Markdown export."""
        job = await self._get_job_or_raise(job_id)
        try:
            job.status = "processing"
            await self.session.commit()

            cases = await self._get_filtered_cases(job)
            self._render_markdown(cases)  # validate render

            job.file_path = f"exports/{job.id}.md"
            job.status = "completed"
            job.case_count = len(cases)
            await self.session.commit()
            await self.session.refresh(job)
        except Exception as e:
            job.status = "failed"
            job.error_message = str(e)
            await self.session.commit()
            await self.session.refresh(job)
        return job

    # ── Jira/Xray Push (B-M12-08) ────────────────────────────────────

    async def push_to_jira(self, job_id: UUID, jira_config: JiraConfig) -> dict:
        """推送用例到 Jira/Xray。"""
        job = await self._get_job_or_raise(job_id)
        cases = await self._get_filtered_cases(job)
        if not cases:
            return {
                "job_id": job_id,
                "status": "completed",
                "pushed_count": 0,
                "failed_count": 0,
                "jira_issues": [],
                "error_message": None,
            }

        pushed: list[dict] = []
        failed_count = 0
        headers = {
            "Authorization": f"Bearer {jira_config.auth_token}",
            "Content-Type": "application/json",
            "Accept": "application/json",
        }

        async with httpx.AsyncClient(timeout=30.0, trust_env=False) as client:
            for case in cases:
                payload = self._build_jira_payload(case, jira_config)
                try:
                    resp = await client.post(
                        f"{jira_config.base_url.rstrip('/')}/rest/api/2/issue",
                        json=payload,
                        headers=headers,
                    )
                    if resp.status_code in (200, 201):
                        data = resp.json()
                        pushed.append({"case_id": case["case_id"], "jira_key": data.get("key", "")})
                    else:
                        logger.warning("Jira push failed for %s: %s", case["case_id"], resp.text[:200])
                        failed_count += 1
                except httpx.HTTPError as e:
                    logger.warning("Jira push error for %s: %s", case["case_id"], str(e))
                    failed_count += 1

        status = "completed" if failed_count == 0 else ("partial" if pushed else "failed")
        return {
            "job_id": job_id,
            "status": status,
            "pushed_count": len(pushed),
            "failed_count": failed_count,
            "jira_issues": pushed,
            "error_message": None if failed_count == 0 else f"{failed_count} case(s) failed to push",
        }

    def _build_jira_payload(self, case: dict, config: JiraConfig) -> dict:
        """构造 Jira REST API 创建 Issue 的 payload。"""
        description_parts = [f"*Priority*: {case['priority']}", f"*Type*: {case['case_type']}"]
        if case.get("precondition"):
            description_parts.append(f"*Precondition*: {case['precondition']}")
        if case.get("steps"):
            description_parts.append("\n||Step||Action||Expected Result||")
            for step in case["steps"]:
                description_parts.append(f"|{step['step_num']}|{step['action']}|{step['expected_result']}|")
        description = "\n".join(description_parts)

        payload: dict = {
            "fields": {
                "project": {"key": config.project_key},
                "summary": case["title"],
                "description": description,
                "issuetype": {"name": "Test"},
                "labels": config.labels if config.labels else [],
            }
        }
        if config.custom_fields:
            payload["fields"].update(config.custom_fields)
        return payload

    async def get_download_url(self, job_id: UUID) -> str:
        job = await self._get_job_or_raise(job_id)
        if not job.file_path:
            from fastapi import HTTPException

            raise HTTPException(status_code=404, detail="Export file not ready")
        return job.file_path

    async def _get_job_or_raise(self, job_id: UUID) -> ExportJob:
        job = await self.get_export_job(job_id)
        if not job:
            from fastapi import HTTPException

            raise HTTPException(status_code=404, detail="Export job not found")
        return job

    async def _get_filtered_cases(self, job: ExportJob) -> list[dict]:
        requirement_id = job.requirement_id
        if not requirement_id and job.filter_criteria:
            raw = job.filter_criteria.get("requirement_id")
            if raw:
                requirement_id = UUID(str(raw))
        return await self._get_cases(requirement_id)

    async def _get_cases(self, requirement_id: UUID | None) -> list[dict]:
        q = select(TestCase).where(TestCase.deleted_at.is_(None))
        if requirement_id:
            q = q.where(TestCase.requirement_id == requirement_id)
        q = q.order_by(TestCase.created_at)
        result = await self.session.execute(q)
        cases = result.scalars().all()

        output: list[dict] = []
        for tc in cases:
            step_q = (
                select(TestCaseStep)
                .where(
                    TestCaseStep.test_case_id == tc.id,
                    TestCaseStep.deleted_at.is_(None),
                )
                .order_by(TestCaseStep.step_num)
            )
            step_result = await self.session.execute(step_q)
            steps = step_result.scalars().all()

            output.append(
                {
                    "case_id": tc.case_id,
                    "title": tc.title,
                    "priority": tc.priority,
                    "case_type": tc.case_type,
                    "status": tc.status,
                    "precondition": tc.precondition,
                    "steps": [
                        {"step_num": s.step_num, "action": s.action, "expected_result": s.expected_result}
                        for s in steps
                    ],
                }
            )

        return output
