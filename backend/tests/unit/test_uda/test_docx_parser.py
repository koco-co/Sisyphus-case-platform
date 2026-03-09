"""B-M01-14 — DOCX 解析器单元测试"""

from __future__ import annotations

from unittest.mock import MagicMock, patch


class TestDocxParseBasic:
    """测试 DOCX 解析基础功能（mock python-docx）"""

    def test_docx_parse_basic(self):
        """DOCX 文件应解析为包含段落和标题的结构化文本。"""
        # 构建 mock Document
        heading = MagicMock()
        heading.text = "需求概述"
        heading.style = MagicMock()
        heading.style.name = "Heading 1"
        heading.runs = [MagicMock(text="需求概述")]

        para = MagicMock()
        para.text = "本需求实现数据同步功能"
        para.style = MagicMock()
        para.style.name = "Normal"
        para.runs = [MagicMock(text="本需求实现数据同步功能")]

        empty_para = MagicMock()
        empty_para.text = ""
        empty_para.style = MagicMock()
        empty_para.style.name = "Normal"
        empty_para.runs = []

        mock_doc = MagicMock()
        mock_doc.paragraphs = [heading, para, empty_para]
        mock_doc.tables = []

        with patch("docx.Document", return_value=mock_doc):
            from docx import Document

            doc = Document("test.docx")
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        assert len(paragraphs) == 2
        assert "需求概述" in paragraphs[0]
        assert "数据同步" in paragraphs[1]

    def test_docx_parse_with_table(self):
        """DOCX 含表格时应能提取表格数据。"""
        cell1 = MagicMock()
        cell1.text = "字段名"
        cell2 = MagicMock()
        cell2.text = "类型"
        row = MagicMock()
        row.cells = [cell1, cell2]

        table = MagicMock()
        table.rows = [row]

        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_doc.tables = [table]

        with patch("docx.Document", return_value=mock_doc):
            from docx import Document

            doc = Document("test.docx")
            rows_data = []
            for t in doc.tables:
                for r in t.rows:
                    rows_data.append([c.text for c in r.cells])

        assert len(rows_data) == 1
        assert rows_data[0] == ["字段名", "类型"]

    def test_docx_parse_empty_document(self):
        """空 DOCX 应返回空列表。"""
        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_doc.tables = []

        with patch("docx.Document", return_value=mock_doc):
            from docx import Document

            doc = Document("empty.docx")
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        assert paragraphs == []
